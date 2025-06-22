import io
import logging
from typing import Optional, List
from PIL import Image
import pypdf
from pdf2image import convert_from_bytes
from docx import Document as DocxDocument
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import zipfile

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing various document formats"""
    
    def __init__(self):
        self.pdf_dpi = 300  # High DPI for better OCR
        self.max_pages = 5  # Process only first 5 pages to avoid memory issues
    
    async def pdf_to_image(self, pdf_bytes: bytes) -> Optional[bytes]:
        """
        Convert PDF to image for OCR processing
        
        Args:
            pdf_bytes: PDF file content as bytes
            
        Returns:
            Image bytes of the first page with receipt-like content
        """
        try:
            # Try to use pdf2image (requires poppler)
            try:
                images = convert_from_bytes(
                    pdf_bytes, 
                    dpi=self.pdf_dpi,
                    first_page=1,
                    last_page=min(self.max_pages, 5)
                )
                
                # Find the best page (usually first one for receipts)
                for i, image in enumerate(images):
                    logger.info(f"Processing PDF page {i+1}")
                    # Convert PIL image to bytes
                    img_buffer = io.BytesIO()
                    image.save(img_buffer, format='PNG')
                    img_bytes = img_buffer.getvalue()
                    
                    # For receipts, usually the first page is enough
                    if i == 0:
                        return img_bytes
                    
                    # Could add logic here to detect which page has receipt content
                    
                return img_bytes if images else None
                
            except Exception as e:
                logger.warning(f"pdf2image failed (poppler might not be installed): {e}")
                # Fallback to pypdf for text extraction
                return await self._pypdf_fallback(pdf_bytes)
                
        except Exception as e:
            logger.error(f"Error converting PDF to image: {e}", exc_info=True)
            return None
    
    async def _pypdf_fallback(self, pdf_bytes: bytes) -> Optional[bytes]:
        """
        Fallback method using pypdf to extract images from PDF
        """
        try:
            pdf_reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
            
            # Try to extract images from first few pages
            for page_num in range(min(len(pdf_reader.pages), self.max_pages)):
                page = pdf_reader.pages[page_num]
                
                # Extract images from page
                if '/XObject' in page['/Resources']:
                    xObject = page['/Resources']['/XObject'].get_object()
                    
                    for obj in xObject:
                        if xObject[obj]['/Subtype'] == '/Image':
                            # Extract image
                            data = xObject[obj].get_data()
                            
                            # Convert to PIL Image
                            img = Image.open(io.BytesIO(data))
                            
                            # Convert to PNG bytes
                            img_buffer = io.BytesIO()
                            img.save(img_buffer, format='PNG')
                            return img_buffer.getvalue()
            
            # If no images found, create image from text
            text = ""
            for page_num in range(min(len(pdf_reader.pages), self.max_pages)):
                text += pdf_reader.pages[page_num].extract_text()
            
            if text:
                # Create simple image with text for OCR
                return self._text_to_image(text)
                
        except Exception as e:
            logger.error(f"pypdf fallback failed: {e}")
            
        return None
    
    async def extract_images_from_docx(self, docx_bytes: bytes) -> Optional[bytes]:
        """
        Extract images from Word documents
        
        Args:
            docx_bytes: DOCX file content as bytes
            
        Returns:
            First image found in the document
        """
        try:
            # Load document
            doc = DocxDocument(io.BytesIO(docx_bytes))
            
            # Method 1: Check inline shapes
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run._element.xpath('.//a:blip'):
                        # Found an image
                        for blip in run._element.xpath('.//a:blip'):
                            embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if embed:
                                rel = doc.part.rels[embed]
                                if 'image' in rel.target_ref:
                                    image_data = rel.target_part.blob
                                    return image_data
            
            # Method 2: Extract from docx zip
            with zipfile.ZipFile(io.BytesIO(docx_bytes)) as docx_zip:
                # List all files in the zip
                for file_info in docx_zip.filelist:
                    if file_info.filename.startswith('word/media/') and \
                       any(file_info.filename.endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.bmp']):
                        # Extract image
                        image_data = docx_zip.read(file_info.filename)
                        return image_data
            
            # If no images found, extract text and convert to image
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text])
            if text:
                return self._text_to_image(text)
                
        except Exception as e:
            logger.error(f"Error extracting images from DOCX: {e}", exc_info=True)
            
        return None
    
    def _text_to_image(self, text: str, max_chars: int = 2000) -> bytes:
        """
        Convert text to a simple image for OCR processing
        
        Args:
            text: Text content
            max_chars: Maximum characters to include
            
        Returns:
            Image bytes
        """
        try:
            from PIL import Image, ImageDraw, ImageFont
            
            # Limit text length
            text = text[:max_chars]
            
            # Create white background image
            width = 800
            height = max(600, min(2000, len(text) * 2))
            img = Image.new('RGB', (width, height), color='white')
            draw = ImageDraw.Draw(img)
            
            # Try to use a monospace font
            try:
                font = ImageFont.truetype("/System/Library/Fonts/Menlo.ttc", 14)
            except:
                font = ImageFont.load_default()
            
            # Draw text
            margin = 20
            y_position = margin
            
            for line in text.split('\n'):
                if y_position > height - margin:
                    break
                    
                # Wrap long lines
                if len(line) > 80:
                    words = line.split()
                    current_line = ""
                    for word in words:
                        if len(current_line + word) < 80:
                            current_line += word + " "
                        else:
                            draw.text((margin, y_position), current_line.strip(), font=font, fill='black')
                            y_position += 20
                            current_line = word + " "
                    if current_line:
                        draw.text((margin, y_position), current_line.strip(), font=font, fill='black')
                        y_position += 20
                else:
                    draw.text((margin, y_position), line, font=font, fill='black')
                    y_position += 20
            
            # Convert to bytes
            img_buffer = io.BytesIO()
            img.save(img_buffer, format='PNG')
            return img_buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Error converting text to image: {e}")
            # Return a minimal image
            img = Image.new('RGB', (100, 100), color='white')
            img_buffer = io.BytesIO()
            img.save(img_buffer, format='PNG')
            return img_buffer.getvalue()
    
    async def validate_file_type(self, file_bytes: bytes, expected_mime: str) -> bool:
        """
        Validate file type using python-magic
        
        Args:
            file_bytes: File content
            expected_mime: Expected MIME type
            
        Returns:
            True if file type matches
        """
        try:
            import magic
            mime = magic.from_buffer(file_bytes, mime=True)
            return mime == expected_mime
        except:
            # If python-magic not available, trust the provided MIME type
            return True